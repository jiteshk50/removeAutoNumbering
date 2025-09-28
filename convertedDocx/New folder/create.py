# requirements:
#   pip install python-docx

import re
import os
from docx import Document
from docx.shared import Inches

# Configuration
INPUT_DIR = "."              # Folder containing set_1.docx, set_2.docx, ...
FILE_PATTERN = "set_{}.docx" # Name pattern
NUM_FILES = 10               # set_1 ... set_10
NUM_QUESTIONS = 12           # Q1 ... Q12 expected
OUTPUT_DIR = "."             # Where to write Q1.docx ... Q12.docx

# Output formatting
LEFT_INDENT_IN = 0.25        # indent for the whole first line ("n. text...")
SPACE_BETWEEN_ENTRIES = True

# Helper: extract question blocks by numeric numbering at paragraph start.
# Assumes each question starts with a paragraph beginning with "<n>" optionally followed by '.' or ')'
def extract_questions_from_docx(path, max_q=12):
    doc = Document(path)
    header_re = re.compile(r"^\s*(\d{1,3})\s*[\.\)]?\s*(.*)$")
    paras = doc.paragraphs

    # Identify candidate starts
    starts = []
    for i, p in enumerate(paras):
        m = header_re.match(p.text.strip())
        if m:
            try:
                num = int(m.group(1))
                if 1 <= num <= max_q:
                    starts.append((num, i))
            except ValueError:
                pass

    # Keep first occurrence per question number
    first_occurrence = {}
    for num, idx in starts:
        if num not in first_occurrence:
            first_occurrence[num] = idx

    ordered = sorted([(n, i) for n, i in first_occurrence.items()], key=lambda x: x[1])
    ranges = []
    for k, (num, start_idx) in enumerate(ordered):
        end_idx = ordered[k + 1][1] if k + 1 < len(ordered) else len(paras)
        ranges.append((num, start_idx, end_idx))

    # Build num -> list of paragraph texts for that question
    questions = {}
    for num, s, e in ranges:
        block = paras[s:e]
        cleaned = []
        first = True
        header_re2 = re.compile(r"^\s*(\d{1,3})\s*[\.\)]?\s*(.*)$")
        for p in block:
            text = p.text
            if first:
                m = header_re2.match(text.strip())
                if m and int(m.group(1)) == num:
                    rest = m.group(2).strip()
                    text = rest if rest else ""
                first = False
            cleaned.append(text)
        questions[num] = cleaned

    return questions

def main():
    # Aggregate blocks by question number across files
    aggregated = {q: [] for q in range(1, NUM_QUESTIONS + 1)}

    for i in range(1, NUM_FILES + 1):
        fname = FILE_PATTERN.format(i)
        path = os.path.join(INPUT_DIR, fname)
        if not os.path.exists(path):
            print(f"Warning: missing {path}, skipping")
            continue
        qs = extract_questions_from_docx(path, max_q=NUM_QUESTIONS)
        for qn in range(1, NUM_QUESTIONS + 1):
            block = qs.get(qn)
            if block is not None:
                aggregated[qn].append((fname, block))

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # For each question number, write Q{n}.docx with sequential numbering per entry: "1. text..."
    for qn in range(1, NUM_QUESTIONS + 1):
        out_doc = Document()
        seq = 0
        for _src_name, lines in aggregated[qn]:
            seq += 1
            # Combine number and first line into a single paragraph to avoid a break after the number
            first_line = lines[0] if lines else ""
            para = out_doc.add_paragraph(f"{seq}. {first_line}".rstrip())
            para.paragraph_format.left_indent = Inches(LEFT_INDENT_IN)

            # Add remaining lines as separate paragraphs (unchanged)
            for line in lines[1:]:
                out_doc.add_paragraph(line)

            if SPACE_BETWEEN_ENTRIES:
                out_doc.add_paragraph("")

        out_name = os.path.join(OUTPUT_DIR, f"Q{qn}.docx")
        out_doc.save(out_name)
        print(f"Wrote {out_name} with {seq} entries")

if __name__ == "__main__":
    main()
