import os
from docx import Document
import glob

def get_table_text(table):
    text = []
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            text.append(" ".join(row_text))
    return "\n".join(text)

def extract_question(doc_path, question_number):
    try:
        doc = Document(doc_path)
        content = []
        question_content = []
        capturing = False
        next_question_number = question_number + 1

        # First, create a combined list of content from both paragraphs and tables
        for element in doc.element.body:
            if element.tag.endswith('p'):  # It's a paragraph
                content.append(("p", doc.paragraphs[len([p for p in content if p[0] == "p"])].text))
            elif element.tag.endswith('tbl'):  # It's a table
                table_idx = len([t for t in content if t[0] == "tbl"])
                table_text = get_table_text(doc.tables[table_idx])
                if table_text.strip():
                    content.append(("tbl", table_text))

        # Now process the combined content
        for content_type, text in content:
            text = text.strip()
            if not text:
                continue

            # Start capturing when we find our target question
            if text.startswith(f"{question_number}."):
                capturing = True
                question_content.append(text)
                continue
            
            # Stop capturing if we hit the next question number
            if capturing and text.startswith(f"{next_question_number}."):
                break
            
            # While capturing, append non-empty content
            if capturing and text:
                question_content.append(text)
            
        return "\n".join(question_content) if question_content else None
    except Exception as e:
        print(f"Error reading {doc_path}: {str(e)}")
    return None

def create_question_paper():
    # Get user input for question number
    while True:
        try:
            question_number = int(input("Enter the question number to extract (1-20): "))
            if 1 <= question_number <= 20:
                break
            else:
                print("Please enter a number between 1 and 20")
        except ValueError:
            print("Please enter a valid number")

    # Create set directory if it doesn't exist
    set_dir = "set"
    if not os.path.exists(set_dir):
        os.makedirs(set_dir)

    # Create a new document
    new_doc = Document()
    new_doc.add_heading(f'Question Paper - Questions #{question_number} from all files', 0)

    # Get all 100 question files
    docx_files = []
    for i in range(1, 101):
        docx_files.append(f"Q{i}.docx")

    question_count = 0
    missing_files = []
    error_files = []
    processed_files = []
    
    for file in docx_files:
        if not os.path.exists(file):
            missing_files.append(file)
            continue
            
        try:
            question = extract_question(file, question_number)
            if question:
                question_count += 1
                processed_files.append(file)
                # Add the complete question with its options
                new_doc.add_paragraph(f"{question_count}. {question[len(str(question_number))+2:]}")
                # Add a blank line between questions for better readability
                new_doc.add_paragraph()
            else:
                error_files.append((file, "No question found"))
        except Exception as e:
            error_files.append((file, str(e)))
    
    # Print debug information
    print("\n=== Processing Summary ===")
    print(f"Total files processed: {len(processed_files)}")
    print(f"Missing files: {len(missing_files)}")
    if missing_files:
        print("Missing files:", ", ".join(missing_files[:5]) + ("..." if len(missing_files) > 5 else ""))
    print(f"Files with errors: {len(error_files)}")
    if error_files:
        print("First few errors:")
        for file, error in error_files[:3]:
            print(f"  {file}: {error}")
    print("=======================\n")

    # Save the new document in the set folder with the new naming format
    output_filename = os.path.join(set_dir, f"set_{question_number}.docx")
    new_doc.save(output_filename)
    print(f"\nCreated {output_filename} with {question_count} questions")

if __name__ == "__main__":
    create_question_paper()